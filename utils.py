import os
import numpy as np
from langchain.tools import Tool
from langchain.agents import initialize_agent, AgentType
from langchain_community.utilities import SQLDatabase
from langchain_community.agent_toolkits.sql.base import create_sql_agent
from langchain.chains import RetrievalQA
from langchain.schema import Document, BaseRetriever
from langchain.embeddings import HuggingFaceEmbeddings
from langchain.vectorstores import FAISS as LangchainFAISS
from docx import Document as DocxDocument
from openpyxl import Workbook
from reportlab.platypus import SimpleDocTemplate, Paragraph
from reportlab.lib.styles import getSampleStyleSheet
from langchain.schema import SystemMessage
import json
import re
import time
import hashlib
import streamlit as st
from dotenv import load_dotenv
import requests
import uuid
from langchain.memory import ConversationBufferMemory
from chat_history import load_all_conversations
from prompt import SYSTEM_PROMPT
import torch
import socket
from typing import Dict, Any, List
import pathlib
import threading
from sqlalchemy import create_engine, inspect
# import pymysql
secret_path = "/etc/secrets/OPENAI_API_KEY"
if os.path.exists(secret_path):
    with open(secret_path) as f:
        os.environ["OPENAI_API_KEY"] = f.read().strip()


if not os.getenv("OPENAI_API_KEY"):
    st.warning("OPENAI_API_KEY is missing at runtime. LLM features will be disabled until configured.")
# Load environment variables from .env file
# load_dotenv()

# Access MySQL credentials
# MYSQL_HOST = os.getenv("MYSQL_HOST")
# MYSQL_USER = os.getenv("MYSQL_USER")
# MYSQL_PASSWORD = os.getenv("MYSQL_PASSWORD")
# MYSQL_DATABASE = os.getenv("MYSQL_DATABASE")
# MYSQL_PORT = os.getenv("MYSQL_PORT", "3306")


MYSQL_HOST = st.secrets["mysql"]["MYSQL_HOST"]
MYSQL_USER = st.secrets["mysql"]["MYSQL_USER"]
MYSQL_PASSWORD = st.secrets["mysql"]["MYSQL_PASSWORD"]
MYSQL_DATABASE = st.secrets["mysql"]["MYSQL_DATABASE"]
MYSQL_PORT = st.secrets["mysql"]["MYSQL_PORT"]
# MYSQL_HOST = os.environ.get("MYSQL_HOST")
# MYSQL_PORT = int(os.environ.get("MYSQL_PORT", 3306))
# MYSQL_USER = os.environ.get("MYSQL_USER")
# MYSQL_PASSWORD = os.environ.get("MYSQL_PASSWORD")
# MYSQL_DATABASE = os.environ.get("MYSQL_DATABASE")


@st.cache_resource
def get_llm(model: str = "gpt-4.1-nano", temperature: float = 0):
    # lazy import ChatOpenAI to avoid heavy import cost at module load
    from langchain_openai import ChatOpenAI
    global _LLM_SINGLETON
    try:
        if _LLM_SINGLETON is not None:
            return _LLM_SINGLETON
    except NameError:
        _LLM_SINGLETON = None

    if not os.getenv("OPENAI_API_KEY"):
        raise RuntimeError("OPENAI_API_KEY not set; call get_llm() after configuring the key")

    _LLM_SINGLETON = ChatOpenAI(model=model, temperature=temperature,max_tokens=512)
    return _LLM_SINGLETON


# Global LLM singleton (initialized lazily)
_LLM_SINGLETON = None

# ---------- Index & embedding cache settings ----------
INDEX_DIR = os.path.join("data", ".index")
INDEX_FAISS_PATH = os.path.join(INDEX_DIR, "index.faiss")
INDEX_DOCS_PATH = os.path.join(INDEX_DIR, "docs.json")
INDEX_META_PATH = os.path.join(INDEX_DIR, "meta.json")
INDEX_HASHES_PATH = os.path.join(INDEX_DIR, "file_hashes.json")

# Singleton embedding wrapper (LangChain HuggingFaceEmbeddings)
_EMBEDDING_WRAPPER = None

def get_embedding_wrapper(model_name: str = 'sentence-transformers/all-MiniLM-L6-v2'):
    global _EMBEDDING_WRAPPER
    if _EMBEDDING_WRAPPER is None:
        _EMBEDDING_WRAPPER = get_cpu_huggingface_embeddings(model_name)
    return _EMBEDDING_WRAPPER

def _compute_file_hashes(data_folder: str = "data") -> Dict[str, Dict[str, int]]:
    """
    Return a fast signature for files in `data_folder` keyed by filename -> {mtime, size}.
    This avoids reading entire file contents and is much faster for large files.
    """
    sigs = {}
    try:
        for file in os.listdir(data_folder):
            file_path = os.path.join(data_folder, file)
            if os.path.isfile(file_path):
                try:
                    stt = os.stat(file_path)
                    sigs[file] = {"mtime": int(stt.st_mtime), "size": stt.st_size}
                except Exception:
                    continue
    except Exception:
        pass
    return sigs


def _safe_name(filename: str) -> str:
    return hashlib.sha1(filename.encode("utf-8")).hexdigest()


def _per_file_paths(filename: str):
    safe = _safe_name(filename)
    base = os.path.join(INDEX_DIR, safe)
    pathlib.Path(base).mkdir(parents=True, exist_ok=True)
    return {
        "dir": base,
        "docs": os.path.join(base, "docs.json"),
        "meta": os.path.join(base, "meta.json"),
        "emb": os.path.join(base, "emb.npy"),
        "hash": os.path.join(base, "hash.txt"),
    }


# Background-threaded indexing removed — rely on `@st.cache_resource` for cached index builds


# Cached global index wrapper using file-hash as cache key
@st.cache_resource
def cached_build_index(file_sig: str, data_folder="data", chunk_size: int = 300, overlap: int = 20, batch_size: int = 64):
    # file_sig should be a deterministic but cheap fingerprint (mtime+size) of files
    return build_index(data_folder=data_folder, chunk_size=chunk_size, overlap=overlap, batch_size=batch_size)



# -------------------------------
# SentenceTransformer helpers
# -------------------------------
def get_cpu_huggingface_embeddings(
    model_name: str = "sentence-transformers/all-MiniLM-L6-v2",
):
    """LangChain wrapper – always safe on CPU."""
    from langchain.embeddings import HuggingFaceEmbeddings

    return HuggingFaceEmbeddings(
        model_name="sentence-transformers/all-MiniLM-L6-v2",
        model_kwargs={"device": "cpu"},
        encode_kwargs={"normalize_embeddings": True},
    )


def get_cpu_sentence_transformer(model_name: str = "all-MiniLM-L6-v2"):
    """SentenceTransformer – never calls .to() on a meta tensor."""
    from sentence_transformers import SentenceTransformer
    import torch

    model = SentenceTransformer(model_name)               # load without device
    if any(p.device.type == "meta" for p in model.parameters()):
        model = SentenceTransformer(model_name, device="cpu")
    else:
        model = model.to(torch.device("cpu"))
    return model

def extract_text_from_pdf(pdf_path):
    # import here to avoid heavy import at module load
    import fitz
    text = ""
    with fitz.open(pdf_path) as doc:
        for page in doc:
            page_text = page.get_text()
            if page_text:
                text += page_text + "\n"
    return text

def extract_structured_from_excel(excel_path):
    import pandas as pd
    df = pd.read_excel(excel_path)
    df = df.fillna("")
    structured_rows = []
    text_rows = []
    for _, row in df.iterrows():
        row_dict = {col: str(row[col]) for col in df.columns}
        row_text = " | ".join([f"{col}: {row_dict[col]}" for col in df.columns])
        structured_rows.append(row_dict)
        text_rows.append(row_text)
    return structured_rows, "\n".join(text_rows)

def chunk_text(text, chunk_size=300, overlap=50):
    words = text.split()
    chunks = []
    start = 0
    while start < len(words):
        end = start + chunk_size
        chunk = " ".join(words[start:end])
        chunks.append(chunk)
        start += chunk_size - overlap
    return chunks

def build_index(data_folder="data", chunk_size: int = 300, overlap: int = 20, batch_size: int = 64):
    docs, metadata = [], []

    pathlib.Path(INDEX_DIR).mkdir(parents=True, exist_ok=True)
    # lazy import faiss to avoid import cost at module load
    import faiss

    # Quick-check: if a saved global index exists and files haven't changed, load it
    try:
        current_hashes = _compute_file_hashes(data_folder)
        if os.path.exists(INDEX_FAISS_PATH) and os.path.exists(INDEX_DOCS_PATH) and os.path.exists(INDEX_META_PATH) and os.path.exists(INDEX_HASHES_PATH):
            try:
                with open(INDEX_HASHES_PATH, "r", encoding="utf-8") as fh:
                    saved_hashes = json.load(fh)
                if saved_hashes == current_hashes:
                    # load persisted index & metadata
                    with open(INDEX_DOCS_PATH, "r", encoding="utf-8") as fh:
                        docs = json.load(fh)
                    with open(INDEX_META_PATH, "r", encoding="utf-8") as fh:
                        metadata = json.load(fh)
                    index = faiss.read_index(INDEX_FAISS_PATH)
                    return docs, metadata, index
            except Exception:
                # If any load step fails, fall back to rebuilding
                pass
    except Exception:
        # ignore hashing errors and rebuild
        pass

    # We'll build per-file embeddings (cached) and then combine
    embedder = get_embedding_wrapper()
    all_embs = []
    doc_list = []
    meta_list = []

    files = [f for f in os.listdir(data_folder) if os.path.isfile(os.path.join(data_folder, f))]
    total_files = len(files)
    file_idx = 0

    for file in files:
        file_idx += 1
        file_path = os.path.join(data_folder, file)
        paths = _per_file_paths(file)

        # compute a cheap file signature (mtime:size) and compare with cached signature
        try:
            stt = os.stat(file_path)
            file_sig = f"{int(stt.st_mtime)}:{stt.st_size}"
        except Exception:
            file_sig = None

        use_cache = False
        if file_sig and os.path.exists(paths["hash"]):
            try:
                with open(paths["hash"], "r", encoding="utf-8") as fh:
                    saved = fh.read().strip()
                if saved == file_sig and os.path.exists(paths["emb"]) and os.path.exists(paths["docs"]) and os.path.exists(paths["meta"]):
                    use_cache = True
            except Exception:
                use_cache = False

        if use_cache:
            try:
                file_docs = json.load(open(paths["docs"], "r", encoding="utf-8"))
                file_meta = json.load(open(paths["meta"], "r", encoding="utf-8"))
                file_emb = np.load(paths["emb"])
            except Exception:
                use_cache = False
        print(use_cache)
        if not use_cache:
            # extract text and chunk depending on filetype
            if file.lower().endswith(".pdf"):
                print("i am pdf")
                text = extract_text_from_pdf(file_path)
                chunks = chunk_text(text, chunk_size=chunk_size, overlap=overlap)
                file_docs = [c.strip() for c in chunks if c.strip()]
                file_meta = [{"source": file, "type": "pdf", "structured": None} for _ in file_docs]
            elif file.lower().endswith((".xlsx", ".xls")):
                structured_rows, text_data = extract_structured_from_excel(file_path)
                chunks = chunk_text(text_data, chunk_size=chunk_size, overlap=overlap)
                file_docs = [c.strip() for c in chunks if c.strip()]
                file_meta = [{"source": file, "type": "excel", "structured": structured_rows} for _ in file_docs]
            elif file.lower().endswith((".txt", ".md")):
                try:
                    with open(file_path, "r", encoding="utf-8") as f:
                        text = f.read()
                except Exception:
                    continue
                chunks = chunk_text(text, chunk_size=chunk_size, overlap=overlap)
                file_docs = [c.strip() for c in chunks if c.strip()]
                file_meta = [{"source": file, "type": "text", "structured": None} for _ in file_docs]
            else:
                # unsupported file types are skipped
                continue

            # compute embeddings for this file in batches
            emb_list = []
            for i in range(0, len(file_docs), batch_size):
                batch = file_docs[i : i + batch_size]
                try:
                    batch_emb = embedder.embed_documents(batch)
                except Exception:
                    batch_emb = []
                    for d in batch:
                        batch_emb.append(embedder.embed_documents([d])[0])
                emb_list.extend(batch_emb)

            file_emb = np.array(emb_list, dtype=np.float32)

            # persist per-file cache (store cheap file signature)
            try:
                np.save(paths["emb"], file_emb)
                with open(paths["docs"], "w", encoding="utf-8") as fh:
                    json.dump(file_docs, fh)
                with open(paths["meta"], "w", encoding="utf-8") as fh:
                    json.dump(file_meta, fh)
                if file_sig:
                    with open(paths["hash"], "w", encoding="utf-8") as fh:
                        fh.write(file_sig)
            except Exception:
                pass

        # append to global lists
        if len(file_docs) > 0:
            start_idx = len(doc_list)
            doc_list.extend(file_docs)
            meta_list.extend(file_meta)
            all_embs.append(file_emb)

        # update quick progress in session_state if present
        try:
            st.session_state["index_progress"] = file_idx / max(1, total_files)
        except Exception:
            pass

    if not doc_list:
        dim = 384
        index = faiss.IndexFlatIP(dim)
        return [], [], index

    # concatenate embeddings
    embeddings = np.vstack(all_embs)

    dim = embeddings.shape[1]
    index = faiss.IndexFlatIP(dim)
    index.add(embeddings)

    # Persist global index and metadata for fast reloads
    try:
        faiss.write_index(index, INDEX_FAISS_PATH)
        with open(INDEX_DOCS_PATH, "w", encoding="utf-8") as fh:
            json.dump(doc_list, fh)
        with open(INDEX_META_PATH, "w", encoding="utf-8") as fh:
            json.dump(meta_list, fh)
        with open(INDEX_HASHES_PATH, "w", encoding="utf-8") as fh:
            json.dump(_compute_file_hashes(data_folder), fh)
    except Exception:
        pass

    return doc_list, meta_list, index

def get_sqlite_tool(memory=None):
    db = SQLDatabase.from_uri("sqlite:///structured_data.db")
    system_message = SystemMessage(content=SYSTEM_PROMPT)
    sql_agent_executor = create_sql_agent(
        llm=get_llm(),
        db=db,
        agent_type=AgentType.OPENAI_FUNCTIONS,
        memory=memory,
        verbose=True,
        agent_kwargs={"system_message": system_message}
    )

    def _sqlite_tool_wrapper(q):
        q = normalize_query(q)
        result = sql_agent_executor.invoke({"input": q})
        return result.get("output", "")

    return Tool(
        name="sqlite_tool",
        func=_sqlite_tool_wrapper,
        description="Answer questions related to user's chat history stored in SQLite."
    )


def generate_table_info(engine):
    inspector = inspect(engine)
    table_info = {}

    for table_name in inspector.get_table_names():
        columns = inspector.get_columns(table_name)
        column_lines = [f"    - {col['name']} ({col['type']})" for col in columns]
        description = f"### {table_name}\n- Fields:\n" + "\n".join(column_lines)
        table_info[table_name] = description

    return table_info

def get_mysql_tool(memory=None, db_uri=None):
    if db_uri is None:
        db_uri = f"mysql+pymysql://{MYSQL_USER}:{MYSQL_PASSWORD}@{MYSQL_HOST}:{MYSQL_PORT}/{MYSQL_DATABASE}"

    # Every table referenced anywhere in prompt.py's business rules must be
    # listed here. Anything not on this list is invisible to the SQL agent
    # -- it will report a real, existing table as "does not exist" if it's
    # missing from ALLOWED_TABLES. Add a table here the moment a new
    # business rule in prompt.py references it.
    ALLOWED_TABLES = [
        "tbl_invoices",
        "tbl_invoice_products_details",
        "tbl_purchase_order",
        "tbl_stock_products",
        "tbl_accounting_return_invoices",
        "tbl_common_currency",
        "tbl_common_current_year",
        "tbl_inventory_purchases",
        "tbl_inventory_purchase_items",
        "tbl_uom_conversions",
        "tbl_stock_uom",
    ]

    def _call_mysql_tool(query: str) -> str:
        try:
            # lazy imports
            from sqlalchemy import create_engine, inspect
            from langchain_community.utilities import SQLDatabase
            from langchain_community.agent_toolkits.sql.base import create_sql_agent
            from langchain.schema import SystemMessage

            engine = create_engine(db_uri, connect_args={"connect_timeout": 5})
            inspector = inspect(engine)
            existing_tables = set(inspector.get_table_names())

            valid_tables = [t for t in ALLOWED_TABLES if t in existing_tables]

            # If something on the list isn't actually in the live DB, don't
            # fail silently -- log it so a renamed/dropped table gets caught
            # immediately instead of surfacing later as a confusing answer.
            missing = [t for t in ALLOWED_TABLES if t not in existing_tables]
            if missing:
                print(f"[mysql_tool] WARNING: expected tables not found in live DB: {missing}")

            # No custom_table_info here on purpose: SQLDatabase will
            # introspect the live database directly, so the agent always
            # sees the real, current columns for every included table.
            # The business logic that isn't obvious from column names alone
            # (join keys, is_sales/is_quotation semantics, the revenue
            # formula, etc.) already lives in prompt.py's business rules --
            # duplicating a second, hand-maintained copy here is exactly
            # what drifted out of sync before.
            db = SQLDatabase.from_uri(
                db_uri,
                include_tables=valid_tables,
                sample_rows_in_table_info=3,
            )

            system_message = SystemMessage(content=SYSTEM_PROMPT)
            sql_agent_executor = create_sql_agent(
                llm=get_llm(),
                db=db,
                agent_type=AgentType.OPENAI_FUNCTIONS,
                memory=memory,
                verbose=True,
                agent_kwargs={"system_message": system_message}
            )

            q = normalize_query(query)
            result = sql_agent_executor.invoke({"input": q})
            return result.get("output", "")
        except Exception as e:
            return f"MySQL tool error: {e}"

    return Tool(
        name="mysql_tool",
        func=_call_mysql_tool,
        description=(
            "Answer questions about live ERP data in MySQL: invoices, "
            "purchase orders, purchase invoices, products, currency, and "
            "UOM conversions."
        ),
    )

    def _call_mysql_tool(query: str) -> str:
        try:
            # lazy imports
            from sqlalchemy import create_engine, inspect
            from langchain_community.utilities import SQLDatabase
            from langchain_community.agent_toolkits.sql.base import create_sql_agent
            from langchain.schema import SystemMessage

            engine = create_engine(db_uri, connect_args={"connect_timeout": 5})
            inspector = inspect(engine)
            existing_tables = inspector.get_table_names()

            # Only keep tables that exist in both the DB and your table_info
            valid_tables = [t for t in table_info.keys() if t in existing_tables]

            db = SQLDatabase.from_uri(
                db_uri,
                include_tables=valid_tables,
                sample_rows_in_table_info=3,
                custom_table_info=table_info
            )

            system_message = SystemMessage(content=SYSTEM_PROMPT)
            sql_agent_executor = create_sql_agent(
                llm=get_llm(),
                db=db,
                agent_type=AgentType.OPENAI_FUNCTIONS,
                memory=memory,
                verbose=True,
                agent_kwargs={"system_message": system_message}
            )

            q = normalize_query(query)
            result = sql_agent_executor.invoke({"input": q})
            return result.get("output", "")
        except Exception as e:
            return f"MySQL tool error: {e}"

    return Tool(
        name="mysql_tool",
        func=_call_mysql_tool,
        description="Answer questions about structured financial data stored in MySQL (via phpMyAdmin). Tool creates DB connection lazily when invoked."
    )


def list_mysql_tables():
    db_uri = f"mysql+pymysql://{MYSQL_USER}:{MYSQL_PASSWORD}@{MYSQL_HOST}:{MYSQL_PORT}/{MYSQL_DATABASE}"
    from sqlalchemy import create_engine, inspect
    engine = create_engine(db_uri)
    inspector = inspect(engine)
    print("Tables in DB:", inspector.get_table_names())

def get_retriever_tool(docs, metadata, memory=None):
    print("hi retriever tool called")
    # Use a local retriever that reuses a persisted FAISS index when available
    class LocalRetriever(BaseRetriever):
        index: any
        docs: list
        doc_metadata: list  # Renamed from 'metadata' to avoid shadowing parent class
        embedder: any
        k: int = 4

        def get_relevant_documents(self, query):
            query = normalize_query(query)
            try:
                q_emb = self.embedder.embed_query(query)
            except Exception:
                q_emb = self.embedder.embed_documents([query])[0]

            q_emb = np.array(q_emb, dtype=np.float32)
            if q_emb.ndim == 1:
                q_emb = q_emb.reshape(1, -1)

            D, I = self.index.search(q_emb, self.k)
            results = []

            for idx in I[0]:
                if idx < len(self.docs):
                    # Ensure metadata is a dict with only strings
                    meta = {k: str(v) if v is not None else "" for k, v in self.doc_metadata[idx].items()}
                    results.append(Document(page_content=self.docs[idx], metadata=meta))
            return results

    
    # Build or load index (fast when persisted). Use cached global index when possible.
    file_hashes = _compute_file_hashes("data")
    file_sig = json.dumps(file_hashes, sort_keys=True)

    docs_cached, meta_cached, index = build_index("data")
    print("📦 Index diagnostics")
    print("Docs:", len(docs_cached))
    print("Meta:", len(meta_cached))
    print("Index ntotal:", index.ntotal if index else None)

    # If a session memory is provided, do NOT cache QA chain (avoid caching with memory)
    # But we still don't pass memory to RetrievalQA - let the outer agent handle memory
    if memory is not None:
        import faiss
        embedding = get_embedding_wrapper()
        retr = LocalRetriever(index=index, docs=docs_cached, doc_metadata=meta_cached, embedder=embedding)
        
        def document_retriever_tool(query: str) -> str:
            query = normalize_query(query)
            try:
                # Directly use retriever + LLM, skip RetrievalQA to avoid memory issues
                docs_list = retr.get_relevant_documents(query)
                if not docs_list:
                    return "No relevant documents found."
                
                context = "\n".join([doc.page_content for doc in docs_list])
                
                # Use LLM directly to answer based on retrieved context
                from langchain_core.prompts import PromptTemplate
                prompt = PromptTemplate(
                    template="Based on the following context:\n{context}\n\nAnswer this question: {question}",
                    input_variables=["context", "question"]
                )
                chain = prompt | get_llm()
                result = chain.invoke({"context": context, "question": query})
                return str(result)
            except Exception as e:
                print(f"Error in document_retriever_tool: {e}")
                import traceback
                traceback.print_exc()
                return "⚠️ Could not retrieve answer from documents."

        return Tool(
            name="document_retriever",
            func=document_retriever_tool,
            description="Document retriever (per-session memory)"
        )

    # memory is None -> safe to cache QA chain globally
    cached = st.session_state.get("qa_chain")
    if cached and getattr(cached, "_index_sig", None) == file_sig:
        # cached is a retriever, not a chain
        def document_retriever_tool(query: str) -> str:
            query = normalize_query(query)
            try:
                # Directly use retriever + LLM, skip RetrievalQA to avoid memory issues
                docs_list = cached.get_relevant_documents(query)
                if not docs_list:
                    return "No relevant documents found."
                
                context = "\n".join([doc.page_content for doc in docs_list])
                
                # Use LLM directly to answer based on retrieved context
                from langchain_core.prompts import PromptTemplate
                prompt = PromptTemplate(
                    template="Based on the following context:\n{context}\n\nAnswer this question: {question}",
                    input_variables=["context", "question"]
                )
                chain = prompt | get_llm()
                result = chain.invoke({"context": context, "question": query})
                return str(result)
            except Exception as e:
                print(f"Error in document_retriever_tool: {e}")
                import traceback
                traceback.print_exc()
                return "⚠️ Could not retrieve answer from documents."

        return Tool(
            name="document_retriever",
            func=document_retriever_tool,
            description="Cached document retriever"
        )

    @st.cache_resource
    def _build_retriever(index_hash: str):
        import faiss
        embedding = get_embedding_wrapper()
        retr = LocalRetriever(index=index, docs=docs_cached, doc_metadata=meta_cached, embedder=embedding)
        retr._index_sig = index_hash
        return retr

    retriever = _build_retriever(file_sig)
    
    # Store retriever in session state for potential reuse
    st.session_state["qa_chain"] = retriever

    # 🔑 TOOL CLOSURE — retriever is captured here
    def document_retriever_tool(query: str) -> str:
        query = normalize_query(query)
        try:
            # Directly use retriever + LLM, skip RetrievalQA to avoid memory issues
            docs_list = retriever.get_relevant_documents(query)
            if not docs_list:
                return "No relevant documents found."
            
            context = "\n".join([doc.page_content for doc in docs_list])
            
            # Use LLM directly to answer based on retrieved context
            from langchain_core.prompts import PromptTemplate
            prompt = PromptTemplate(
                template="Based on the following context:\n{context}\n\nAnswer this question: {question}",
                input_variables=["context", "question"]
            )
            chain = prompt | get_llm()
            result = chain.invoke({"context": context, "question": query})
            return str(result)
        except Exception as e:
            print(f"Error in document_retriever_tool: {e}")
            import traceback
            traceback.print_exc()
            return "⚠️ Could not retrieve answer from documents."

    
    st.session_state["qa_chain"] = retriever

    return Tool(
        name="document_retriever",
        func=document_retriever_tool,
        description="Useful for answering questions from resumes, documents, or ERP policies. (cached)"
    )

def get_multi_agent(_, docs, metadata, db_uri=None, memory=None, conversation_history=None):
    if memory is None:
        memory = ConversationBufferMemory(return_messages=True, memory_key="chat_history")

    if conversation_history and not memory.chat_memory.messages:
        for msg in conversation_history:
            if msg["role"] == "user":
                memory.chat_memory.add_user_message(msg["content"])
            elif msg["role"] == "assistant":
                memory.chat_memory.add_ai_message(msg["content"])

    tools = [
        get_sqlite_tool(memory),
        get_retriever_tool(docs, metadata, memory),
    ]

    # Add MySQL tool lazily if host configured (no upfront connection test)
    if MYSQL_HOST:
        if db_uri is None:
            db_uri = (
                f"mysql+pymysql://{MYSQL_USER}:{MYSQL_PASSWORD}"
                f"@{MYSQL_HOST}:{MYSQL_PORT}/{MYSQL_DATABASE}"
            )
        tools.append(get_mysql_tool(memory, db_uri))

    return initialize_agent(
        tools=tools,
        llm=get_llm(),
        agent=AgentType.OPENAI_FUNCTIONS,
        memory=memory,
        verbose=True,
        handle_parsing_errors=True,
        agent_kwargs={"system_message": SystemMessage(content=SYSTEM_PROMPT)}
    )

def generate_chart(query: str, agent=None, retrieved_info: str = "") -> dict:
    default_response = {
        "chart_type": None,
        "chart_data": None,
        "render_method": "plotly",
        "error": "No data available for chart generation"
    }

    # Check if retrieved_info contains valid chart data
    try:
        chart_data = json.loads(retrieved_info) if retrieved_info else {}
        if chart_data.get("labels") and chart_data.get("values") and chart_data.get("title"):
            # Use provided chart data directly for Chart.js
            chartjs_config = {
                "type": "bar",  # Default to bar, adjust based on query if needed
                "data": {
                    "labels": chart_data["labels"],
                    "datasets": [{
                        "label": chart_data.get("title", "Data"),
                        "data": chart_data["values"],
                        "backgroundColor": [
                            "#83c5be", "#006d77", "#ff6b6b", "#4a5568", "#ffd60a",
                            "#8338ec", "#3a86ff", "#f72585"
                        ],
                        "borderColor": "#2d3748",
                        "borderWidth": 1
                    }]
                },
                "options": {
                    "responsive": True,
                    "plugins": {
                        "legend": {"position": "top"},
                        "title": {"display": True, "text": chart_data.get("title", "Chart")}
                    },
                    "scales": {
                        "x": {"title": {"display": True, "text": "Categories"}},
                        "y": {"title": {"display": True, "text": "Values"}}
                    }
                }
            }
            return {
                "chart_type": "bar",
                "chart_data": chartjs_config,
                "render_method": "chartjs",
                "error": None
            }
    except json.JSONDecodeError:
        pass  # Proceed to query the agent if retrieved_info is not valid JSON

    chart_types = {
        "bar": ["bar", "column", "histogram"],
        "line": ["line", "trend"],
        "pie": ["pie", "doughnut"],
        "scatter": ["scatter", "point"],
        "area": ["area"]
    }
    selected_chart_type = None
    for chart_type, keywords in chart_types.items():
        if any(keyword in query.lower() for keyword in keywords):
            selected_chart_type = chart_type
            break
    if not selected_chart_type:
        selected_chart_type = "bar"

    if agent:
        chart_prompt = f"""
        Query the available data sources (preferably get_mysql_tool, then SQL database or document index) to retrieve structured data for the query: '{query}'.
        For product-related queries, always join tbl_invoice_products_details with tbl_stock_products to include Product_Name.
        Return the data in a JSON format suitable for a {selected_chart_type} chart, with:
        - 'labels': list of strings for x-axis or categories
        - 'values': list of numbers for y-axis or data points
        - 'title': string for chart title
        If no relevant data is found, return an empty JSON: {{}}
        """
        try:
            response = agent.invoke({"input": chart_prompt}).get("output", "{}")
            chart_data = json.loads(response)
        except Exception as e:
            return {**default_response, "error": f"Error generating chart data: {e}"}
    else:
        chart_data = {}

    if not chart_data or not chart_data.get("labels") or not chart_data.get("values"):
        return default_response

    # lazy import plotly to avoid heavy import at module load
    import plotly.express as px

    if selected_chart_type == "bar":
        fig = px.bar(x=chart_data["labels"], y=chart_data["values"], title=chart_data.get("title", "Chart"))
    elif selected_chart_type == "line":
        fig = px.line(x=chart_data["labels"], y=chart_data["values"], title=chart_data.get("title", "Chart"))
    elif selected_chart_type == "pie":
        fig = px.pie(names=chart_data["labels"], values=chart_data["values"], title=chart_data.get("title", "Chart"))
    elif selected_chart_type == "scatter":
        fig = px.scatter(x=chart_data["labels"], y=chart_data["values"], title=chart_data.get("title", "Chart"))
    else:
        return {**default_response, "error": f"Unsupported chart type: {selected_chart_type}"}

    chartjs_config = {
        "type": selected_chart_type,
        "data": {
            "labels": chart_data["labels"],
            "datasets": [{
                "label": chart_data.get("title", "Data"),
                "data": chart_data["values"],
                "backgroundColor": [
                    "#83c5be", "#006d77", "#ff6b6b", "#4a5568", "#ffd60a",
                    "#8338ec", "#3a86ff", "#f72585"
                ],
                "borderColor": "#2d3748",
                "borderWidth": 1
            }]
        },
        "options": {
            "responsive": True,
            "plugins": {
                "legend": {"position": "top"},
                "title": {"display": True, "text": chart_data.get("title", "Chart")}
            },
            "scales": {
                "x": {"title": {"display": True, "text": "Categories"}},
                "y": {"title": {"display": True, "text": "Values"}}
            } if selected_chart_type in ["bar", "line", "scatter"] else {}
        }
    }

    return {
        "chart_type": selected_chart_type,
        "chart_data": chartjs_config if "chartjs" in query.lower() else fig,
        "render_method": "chartjs" if "chartjs" in query.lower() else "plotly",
        "error": None
    }
def generate_document(query: str, file_format: str = "pdf", agent=None, retrieved_info: str = "") -> str:
    timestamp = str(int(time.time()))
    query_hash = hashlib.md5(query.encode()).hexdigest()[:8]
    file_path = f"generated_report_{timestamp}_{query_hash}.{file_format.lower()}"

    if file_format.lower() in ["docx", "pdf"]:
        if agent:
            agent_prompt = f"""
            Generate concise, factual content for a {file_format.upper()} document based strictly on: '{query}'.
            Use the get_mysql_tool for real-time ERP data, falling back to other tools only if necessary.
            """
            content = agent.invoke({"input": agent_prompt}).get("output", "Insufficient information available")
        else:
            if not retrieved_info.strip():
                retrieved_info = "No relevant information found in the available data sources."
            content = retrieved_info

        if file_format.lower() == "docx":
            doc = DocxDocument()
            doc.add_heading("Generated Document", 0)
            for paragraph in content.split("\n\n"):
                doc.add_paragraph(paragraph)
            doc.save(file_path)
        else:
            doc = SimpleDocTemplate(file_path)
            styles = getSampleStyleSheet()
            flowables = [Paragraph(p, styles["Normal"]) for p in content.split("\n\n")]
            doc.build(flowables)

    elif file_format.lower() == "xlsx":
        rows = []
        if agent:
            agent_prompt = f"""
            Query the available data sources (preferably get_mysql_tool, then SQL database or document index) to extract structured data for this query:
            '{query}'

            ⚠️ VERY IMPORTANT:
            - Return ONLY structured data (plain CSV or JSON list of dicts)
            - First row = column headers
            - Do NOT include explanations or markdown
            - If no data is found, return exactly: Column1,Column2\nNo data available,,
            """
            result = agent.invoke({"input": agent_prompt})
            raw_output = result.get("output", "").strip()

            try:
                data_list = json.loads(raw_output)
                if isinstance(data_list, list) and len(data_list) > 0:
                    headers = list(data_list[0].keys())
                    rows.append(headers)
                    for row in data_list:
                        rows.append([row.get(h, "") for h in headers])
                else:
                    rows = [["No data available"]]
            except json.JSONDecodeError:
                lines = [line for line in raw_output.split("\n") if line.strip()]
                for line in lines:
                    rows.append([cell.strip() for cell in line.split(",")])

        else:
            rows = [["No data available"]]

        wb = Workbook()
        ws = wb.active
        for r_idx, row in enumerate(rows, start=1):
            for c_idx, value in enumerate(row, start=1):
                ws.cell(row=r_idx, column=c_idx, value=value)
        wb.save(file_path)

    else:
        raise ValueError("Unsupported format. Choose from 'pdf', 'docx', 'xlsx'.")

    return file_path

def determine_file_generation(query: str, conversation_history: str) -> tuple[bool, str | None, bool]:
    prompt = f"""
    You are an AI assistant analyzing a user's query to determine if it requires generating a file (e.g., a report, document, or spreadsheet) or a chart (e.g., bar, line, pie). Consider the query and conversation history to understand the user's intent. The query may not explicitly use words like 'generate', 'export', or 'chart' but may imply a need for a file or chart, such as requesting a report, table, or visual representation. If the query is ambiguous, assume no file or chart generation is needed unless clearly implied by context.

    Query: {query}
    Conversation History: {conversation_history}

    Respond with a valid JSON object (no extra text or code block markers like ```json):
    - "generate_file": boolean indicating if a file should be generated (true/false)
    - "file_format": string ("pdf", "docx", "xlsx") or null if no file is needed
    - "generate_chart": boolean indicating if a chart should be generated (true/false)

    Examples:
    - Query: "Give me a report of sales data" -> {{"generate_file": true, "file_format": "xlsx", "generate_chart": false}}
    - Query: "Create a bar chart of sales data" -> {{"generate_file": false, "file_format": null, "generate_chart": true}}
    - Query: "What is the return policy?" -> {{"generate_file": false, "file_format": null, "generate_chart": false}}
    - Query: "Summarize financials in a document with a chart" -> {{"generate_file": true, "file_format": "docx", "generate_chart": true}}
    - Query: "How many customers are there?" -> {{"generate_file": false, "file_format": null, "generate_chart": false}}
    """
    try:
        response = get_llm().invoke(prompt)
        result = response.content.strip()
        result = re.sub(r'^```json\s*|\s*```$', '', result, flags=re.MULTILINE)
        result = result.strip()
        parsed_result = json.loads(result)
        
        generate_file = parsed_result.get("generate_file", False)
        file_format = parsed_result.get("file_format", None)
        generate_chart = parsed_result.get("generate_chart", False)
        
        if generate_file and file_format not in ["pdf", "docx", "xlsx"]:
            print(f"Invalid file format '{file_format}' in LLM response. Defaulting to no file generation.")
            return False, None, generate_chart
        return generate_file, file_format, generate_chart

    except json.JSONDecodeError as e:
        print(f"JSON parsing error: {e}. Raw response: {result}. Defaulting to no file or chart generation.")
        return False, None, False
    except Exception as e:
        print(f"Error in determine_file_generation: {e}. Defaulting to no file or chart generation.")
        return False, None, False

def cleanup_generated_files(directory="."):
    try:
        for file in os.listdir(directory):
            if file.startswith("generated_report_") and file.endswith((".xlsx", ".pdf", ".docx")):
                file_path = os.path.join(directory, file)
                os.remove(file_path)
                print(f"Deleted file: {file_path}")
    except Exception as e:
        print(f"Error during cleanup of generated files: {e}")

def load_excel_to_db(excel_path, db_uri="sqlite:///structured_data.db"):
    table_name = os.path.splitext(os.path.basename(excel_path))[0]

    import pandas as pd
    from sqlalchemy import create_engine

    df = pd.read_excel(excel_path)
    df = df.fillna("")

    df.columns = (
        df.columns
          .astype(str)
          .str.strip()
          .str.replace(r"\s+", " ", regex=True)
    )

    engine = create_engine(db_uri)
    df.to_sql(table_name, engine, if_exists="replace", index=False)

    return df, table_name


def build_full_user_history(username):
    if not username:
        return []
    return load_all_conversations(username)

def initialize_session_state():
    defaults = {
        "history": [],
        "generated_files": [],
        "generated_charts": [],
        "conversations": {},
        "active_conv": {},
        "authenticated": None,
        "user": None,
        "uploaded_files": [],
        "session_id": str(uuid.uuid4()),
        "memory": None
    }
    for key, value in defaults.items():
        if key not in st.session_state:
            st.session_state[key] = value

    if st.session_state.get("authenticated") and st.session_state.get("user"):
        if st.session_state.get("memory") is None:
            st.session_state.memory = ConversationBufferMemory(
                return_messages=True,
                memory_key="chat_history"
            )
            username = st.session_state.user
            all_history = build_full_user_history(username)
            for msg in all_history:
                if msg["role"] == "user":
                    st.session_state.memory.chat_memory.add_user_message(msg["content"])
                elif msg["role"] == "assistant":
                    st.session_state.memory.chat_memory.add_ai_message(msg["content"])

def normalize_query(q):
    if isinstance(q, dict):
        return q.get("query") or q.get("input") or str(q)
    return str(q)
